import click
from pathlib import Path

import docx
from docx.document import Document


@click.command()
@click.argument('doc_path', type=click.Path(exists=True))
def extract_abbreviations(doc_path: str):
    # Load the document
    doc: Document = docx.Document(doc_path)

    # Dictionary to store abbreviations
    abbreviations = {}
    terms = {}
    prev_run_text = ''

    # Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        try:
            # Check for comments in the paragraph
            for run in paragraph.runs:
                if run.text:
                    prev_run_text = run.text.strip()

                if run.comments:
                    for comment in run.comments:
                        comment_text = comment.text.strip()

                        # Check if comment starts with '@'
                        if comment_text.startswith('@'):
                            abbreviations[prev_run_text] = comment_text[1:]

                        # Check if comment starts with '#'
                        if comment_text.startswith('#'):
                            [term, definition] = comment_text[1:].split(':', 1)
                            terms[term] = definition.strip()
        except Exception as e:
            print(f'Error processing paragraph: {e}')
            print(paragraph.runs[0].text if paragraph.runs else 'empty paragraph')

    print(abbreviations)
    print(terms)

    abbr_file_name = Path(doc_path).name.replace('.docx', '_abbr.docx')
    abbr_doc = docx.Document()
    table = abbr_doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Abbreviation'
    header_cells[1].text = 'Meaning'

    # Add abbreviations to the table
    for abbr, meaning in abbreviations.items():
        row_cells = table.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = meaning

    table = abbr_doc.add_table(rows=1, cols=2)
    # Save the new document
    abbr_doc.save(abbr_file_name)
    print(f'Abbreviations table has been written to {abbr_file_name}')

    # Create and save terms document
    terms_file_name = Path(doc_path).name.replace('.docx', '_terms.docx')
    terms_doc = docx.Document()
    terms_table = terms_doc.add_table(rows=1, cols=2)
    terms_table.style = 'Table Grid'

    # Add header row for terms
    terms_header_cells = terms_table.rows[0].cells
    terms_header_cells[0].text = 'Term'
    terms_header_cells[1].text = 'Definition'

    # Add terms to the table
    for term, definition in terms.items():
        row_cells = terms_table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = definition

    # Save the terms document
    terms_doc.save(terms_file_name)
    print(f'Terms table has been written to {terms_file_name}')


if __name__ == '__main__':
    extract_abbreviations()
