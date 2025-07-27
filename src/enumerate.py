import re
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document


def find_and_replace_patterns(
    input_file: str,
    output_file: str,
    prefixes: List[str] | None = None,
    pattern: str = r'\[([A-Z]+)-[Xx][Xx][Xx]\]'
) -> Dict[str, List[str]]:
    """
    Find and replace text patterns like [REQ-XXX] with sequential numbers.

    Args:
        input_file: Path to input Word document
        output_file: Path to output Word document
        prefixes: List of prefixes to process (e.g., ['REQ', 'SYS'])
        pattern: Regex pattern to match (default: [PREFIX-XXX] case insensitive)

    Returns:
        Dictionary mapping prefixes to lists of generated numbers
    """
    if prefixes is None:
        prefixes = ['REQ', 'SYS']

    # Initialize counters for each prefix
    counters = {prefix: 1 for prefix in prefixes}
    replacements = {prefix: [] for prefix in prefixes}

    try:
        # Load the document
        doc = Document(input_file)

        # First pass: collect all patterns and their locations
        all_matches = _collect_all_patterns(doc, pattern, prefixes)

        # Second pass: process patterns in order for sequential numbering
        _process_all_patterns(doc, all_matches, counters, replacements)

        # Save the modified document
        doc.save(output_file)

        return replacements

    except Exception as e:
        print(f"Error processing document: {e}")
        raise


def _collect_all_patterns(doc, pattern: str, prefixes: List[str]) -> List[Tuple]:
    """
    Collect all patterns and their locations for ordered processing.
    Returns list of (prefix, original_text, location_info) tuples.
    """
    all_matches = []
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        matches = _collect_paragraph_patterns(paragraph, pattern, prefixes)
        all_matches.extend(matches)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = _collect_paragraph_patterns(paragraph, pattern, prefixes)
                    all_matches.extend(matches)

    # Process headers and footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                matches = _collect_paragraph_patterns(paragraph, pattern, prefixes)
                all_matches.extend(matches)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                matches = _collect_paragraph_patterns(paragraph, pattern, prefixes)
                all_matches.extend(matches)

    # Sort matches by prefix for sequential numbering
    all_matches.sort(key=lambda x: x[0])  # Sort by prefix only
    return all_matches


def _collect_paragraph_patterns(paragraph, pattern: str, prefixes: List[str]) -> List[Tuple]:
    """
    Collect patterns from a paragraph including tracked changes.
    """
    matches = []
    
    # Get all runs in the paragraph
    runs = list(paragraph.runs)
    
    if not runs:
        return matches

    # Process each run individually
    for run in runs:
        run_text = run.text
        if not run_text.strip():
            continue
            
        # Find matches in this run
        run_matches = list(re.finditer(pattern, run_text))
        
        for match in run_matches:
            prefix = match.group(1)
            if prefix in prefixes:
                matches.append((
                    prefix,
                    match.group(0),
                    (paragraph, run, match.start(), match.end())
                ))

    # Also collect from XML structure for tracked changes
    xml_matches = _collect_xml_patterns(paragraph, pattern, prefixes)
    matches.extend(xml_matches)
    
    return matches


def _collect_xml_patterns(paragraph, pattern: str, prefixes: List[str]) -> List[Tuple]:
    """
    Collect patterns from XML structure (tracked changes).
    """
    matches = []
    
    try:
        # Get the paragraph's XML element
        p_element = paragraph._element
        
        # Process insertion elements (w:ins)
        for ins_elem in p_element.findall('.//w:ins', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            xml_matches = _collect_xml_element_patterns(ins_elem, pattern, prefixes, paragraph)
            matches.extend(xml_matches)
        
        # Process deletion elements (w:del)
        for del_elem in p_element.findall('.//w:del', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            xml_matches = _collect_xml_element_patterns(del_elem, pattern, prefixes, paragraph)
            matches.extend(xml_matches)
            
    except Exception:
        pass
    
    return matches


def _collect_xml_element_patterns(element, pattern: str, prefixes: List[str], paragraph) -> List[Tuple]:
    """
    Collect patterns from XML element for tracked changes.
    """
    matches = []
    
    try:
        # Get all text runs in the element
        for run_elem in element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            text = run_elem.text
            if not text:
                continue
                
            # Find matches in this text
            run_matches = list(re.finditer(pattern, text))
            
            for match in run_matches:
                prefix = match.group(1)
                if prefix in prefixes:
                    matches.append((
                        prefix,
                        match.group(0),
                        (paragraph, run_elem, match.start(), match.end())
                    ))
                
    except Exception:
        pass
    
    return matches


def _process_all_patterns(doc, all_matches: List[Tuple], counters: Dict[str, int], replacements: Dict[str, List[str]]):
    """
    Process all collected patterns in order for sequential numbering.
    """
    for prefix, original_text, location_info in all_matches:
        # Generate replacement
        number = counters[prefix]
        counters[prefix] += 1
        replacement = f'[{prefix}-{number:03d}]'
        replacements[prefix].append(replacement)

        # Replace the pattern at the specified location
        paragraph, element, start_pos, end_pos = location_info
        
        if hasattr(element, 'text'):  # Regular run
            run_text = element.text
            before_text = run_text[:start_pos]
            after_text = run_text[end_pos:]
            element.text = before_text + replacement + after_text
        else:  # XML element
            text = element.text
            before_text = text[:start_pos]
            after_text = text[end_pos:]
            element.text = before_text + replacement + after_text


def _process_document_tracking(
    doc,
    pattern: str,
    prefixes: List[str],
    counters: Dict[str, int],
    replacements: Dict[str, List[str]]
):
    """
    Process tracked changes at the document level.
    """
    try:
        # Get the document's XML element
        doc_element = doc._element
        
        # Process all insertion elements in the document
        for ins_elem in doc_element.findall('.//w:ins', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            _process_xml_element(ins_elem, pattern, prefixes, counters, replacements)
        
        # Process all deletion elements in the document
        for del_elem in doc_element.findall('.//w:del', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            _process_xml_element(del_elem, pattern, prefixes, counters, replacements)
            
    except Exception:
        # If XML processing fails, continue
        pass


def _process_paragraph_with_tracking(
    paragraph,
    pattern: str,
    prefixes: List[str],
    counters: Dict[str, int],
    replacements: Dict[str, List[str]]
):
    """
    Process paragraph including tracked changes (revision marks).
    This method handles both regular text and revision marks.
    """
    # Get all runs in the paragraph
    runs = list(paragraph.runs)
    
    if not runs:
        return

    # Process each run individually to handle tracked changes
    for run in runs:
        run_text = run.text
        if not run_text.strip():
            continue
            
        # Find matches in this run
        matches = list(re.finditer(pattern, run_text))
        
        if not matches:
            continue

        # Process matches in reverse order to maintain positions
        for match in reversed(matches):
            prefix = match.group(1)
            if prefix not in prefixes:
                continue

            # Generate replacement
            number = counters[prefix]
            counters[prefix] += 1
            replacement = f'[{prefix}-{number:03d}]'
            replacements[prefix].append(replacement)

            # Replace the match in this run
            start_pos = match.start()
            end_pos = match.end()
            
            # Split the run text and replace the target
            before_text = run_text[:start_pos]
            after_text = run_text[end_pos:]
            
            # Update the run text
            run.text = before_text + replacement + after_text
            
            # Update run_text for next iteration
            run_text = run.text

    # Also process the paragraph's XML structure for tracked changes
    _process_paragraph_xml_tracking(
        paragraph, pattern, prefixes, counters, replacements
    )


def _process_paragraph_xml_tracking(
    paragraph,
    pattern: str,
    prefixes: List[str],
    counters: Dict[str, int],
    replacements: Dict[str, List[str]]
):
    """
    Process paragraph XML structure to handle tracked changes.
    """
    try:
        # Get the paragraph's XML element
        p_element = paragraph._element
        
        # Process insertion elements (w:ins)
        for ins_elem in p_element.findall('.//w:ins', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            _process_xml_element(ins_elem, pattern, prefixes, counters, replacements)
        
        # Process deletion elements (w:del)
        for del_elem in p_element.findall('.//w:del', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            _process_xml_element(del_elem, pattern, prefixes, counters, replacements)
            
    except Exception:
        # If XML processing fails, fall back to regular processing
        pass


def _process_xml_element(
    element,
    pattern: str,
    prefixes: List[str],
    counters: Dict[str, int],
    replacements: Dict[str, List[str]]
):
    """
    Process XML element for patterns in tracked changes.
    """
    try:
        # Get all text runs in the element
        for run_elem in element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            text = run_elem.text
            if not text:
                continue
                
            # Find matches in this text
            matches = list(re.finditer(pattern, text))
            
            if not matches:
                continue

            # Process matches in reverse order
            for match in reversed(matches):
                prefix = match.group(1)
                if prefix not in prefixes:
                    continue

                # Generate replacement
                number = counters[prefix]
                counters[prefix] += 1
                replacement = f'[{prefix}-{number:03d}]'
                replacements[prefix].append(replacement)

                # Replace the match in the XML text
                start_pos = match.start()
                end_pos = match.end()
                
                before_text = text[:start_pos]
                after_text = text[end_pos:]
                
                # Update the XML text
                run_elem.text = before_text + replacement + after_text
                
                # Update text for next iteration
                text = run_elem.text
                
    except Exception:
        # If XML processing fails, continue
        pass


def _process_paragraph_robust(
    paragraph,
    pattern: str,
    prefixes: List[str],
    counters: Dict[str, int],
    replacements: Dict[str, List[str]]
):
    """
    Process paragraph using a robust approach that preserves formatting.
    This method handles complex document structures and ensures all matches are replaced.
    """
    # Get all text from the paragraph to find matches
    full_text = paragraph.text
    
    # Find all matches in the text
    matches = list(re.finditer(pattern, full_text))
    
    if not matches:
        return

    # Sort matches by position to process in order
    matches.sort(key=lambda m: m.start())
    
    # Process matches in order for sequential numbering
    for match in matches:
        prefix = match.group(1)
        if prefix not in prefixes:
            continue

        # Generate replacement
        number = counters[prefix]
        counters[prefix] += 1
        replacement = f'[{prefix}-{number:03d}]'
        replacements[prefix].append(replacement)

        # Find and replace the specific text in runs
        _replace_in_runs(paragraph, match.group(0), replacement)


def _replace_in_runs(paragraph, old_text: str, new_text: str):
    """
    Replace text in paragraph runs while preserving formatting.
    """
    # Get all runs in the paragraph
    runs = list(paragraph.runs)
    
    if not runs:
        return

    # Find which run contains the text to replace
    current_pos = 0
    target_run_index = -1
    target_start = -1
    target_end = -1
    
    # Locate the run containing the target text
    for i, run in enumerate(runs):
        run_text = run.text
        run_length = len(run_text)
        
        # Check if this run contains the target text
        if old_text in run_text:
            target_run_index = i
            target_start = run_text.find(old_text)
            target_end = target_start + len(old_text)
            break
        
        current_pos += run_length
    
    if target_run_index == -1:
        # Text not found in any single run, try to find across multiple runs
        _replace_across_runs(paragraph, old_text, new_text)
        return
    
    # Replace text in the target run
    target_run = runs[target_run_index]
    run_text = target_run.text
    
    # Split the run text and replace the target
    before_text = run_text[:target_start]
    after_text = run_text[target_end:]
    
    # Update the run text
    target_run.text = before_text + new_text + after_text


def _replace_across_runs(paragraph, old_text: str, new_text: str):
    """
    Handle cases where the target text spans multiple runs.
    """
    # Get the full paragraph text
    full_text = paragraph.text
    
    # Find the position of the old text
    start_pos = full_text.find(old_text)
    if start_pos == -1:
        return
    
    end_pos = start_pos + len(old_text)
    
    # Clear all runs and recreate with replacement
    paragraph.clear()
    
    # Split the text around the replacement
    before_text = full_text[:start_pos]
    after_text = full_text[end_pos:]
    
    # Add the text with replacement
    if before_text:
        paragraph.add_run(before_text)
    
    paragraph.add_run(new_text)
    
    if after_text:
        paragraph.add_run(after_text)


def _process_text(
    text: str,
    pattern: str,
    prefixes: List[str],
    counters: Dict[str, int],
    replacements: Dict[str, List[str]]
) -> str:
    """
    Process text to replace patterns with sequential numbers.

    Args:
        text: Text to process
        pattern: Regex pattern to match
        prefixes: List of valid prefixes
        counters: Current counters for each prefix
        replacements: Dictionary to store replacements

    Returns:
        Processed text with replacements
    """
    def replace_match(match):
        prefix = match.group(1)
        if prefix in prefixes:
            number = counters[prefix]
            counters[prefix] += 1
            replacement = f'[{prefix}-{number:03d}]'
            replacements[prefix].append(replacement)
            return replacement
        return match.group(0)  # Return original if prefix not in list

    return re.sub(pattern, replace_match, text)


def main():
    """Main function for command line usage."""
    import argparse

    parser = argparse.ArgumentParser(
        description='Find and replace text patterns with sequential numbers'
    )
    parser.add_argument('input_file', help='Input Word document path')
    parser.add_argument('output_file', help='Output Word document path')
    parser.add_argument(
        '--prefixes',
        nargs='+',
        default=['REQ', 'SYS'],
        help='Prefixes to process (default: REQ SYS)'
    )
    parser.add_argument(
        '--pattern',
        default=r'\[([A-Z]+)-[Xx][Xx][Xx]\]',
        help='Regex pattern to match (default: [PREFIX-XXX] case insensitive)'
    )

    args = parser.parse_args()

    # Validate input file exists
    if not Path(args.input_file).exists():
        print(f"Error: Input file '{args.input_file}' not found.")
        return 1

    try:
        replacements = find_and_replace_patterns(
            args.input_file,
            args.output_file,
            args.prefixes,
            args.pattern
        )

        print(f"Successfully processed '{args.input_file}' -> '{args.output_file}'")
        for prefix, numbers in replacements.items():
            if numbers:
                print(f"{prefix}: {len(numbers)} replacements")
                print(f"  Examples: {', '.join(numbers[:5])}")
                if len(numbers) > 5:
                    print(f"  ... and {len(numbers) - 5} more")

        return 0

    except Exception as e:
        print(f"Error processing document: {e}")
        return 1


if __name__ == '__main__':
    exit(main())
